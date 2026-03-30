"""
test_ingestion.py - Unit tests for document ingestion modules.
"""

import io
import os
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


# ── PDFExtractor ──────────────────────────────────────────────────────────────

class TestPDFExtractor:
    def test_raises_on_missing_file(self):
        from ingestion.pdf_extractor import PDFExtractor
        extractor = PDFExtractor(ocr_enabled=False)
        with pytest.raises(FileNotFoundError):
            extractor.extract("/nonexistent/path.pdf")

    @patch("ingestion.pdf_extractor.pdfplumber")
    def test_extracts_text_per_page(self, mock_pdfplumber):
        from ingestion.pdf_extractor import PDFExtractor, PageContent

        # Mock pdfplumber page
        mock_page = MagicMock()
        mock_page.page_number = 1
        mock_page.extract_text.return_value = "Sample analysis moisture content 5%"
        mock_page.extract_tables.return_value = []

        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = lambda self: self
        mock_pdf.__exit__ = MagicMock(return_value=False)
        mock_pdfplumber.open.return_value = mock_pdf

        # Use a fake path (file existence check is bypassed by mock)
        with patch("pathlib.Path.exists", return_value=True):
            extractor = PDFExtractor(ocr_enabled=False)
            pages = extractor.extract("/fake/doc.pdf")

        assert len(pages) == 1
        assert "moisture" in pages[0].raw_text.lower()
        assert pages[0].is_ocr is False

    @patch("ingestion.pdf_extractor.pdfplumber")
    def test_extracts_tables(self, mock_pdfplumber):
        from ingestion.pdf_extractor import PDFExtractor

        mock_page = MagicMock()
        mock_page.page_number = 1
        mock_page.extract_text.return_value = "Some text " * 5
        mock_page.extract_tables.return_value = [
            [["Analysis", "Units", "Min", "Max"],
             ["Moisture", "%", "0", "15"]]
        ]

        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page]
        mock_pdf.__enter__ = lambda self: self
        mock_pdf.__exit__ = MagicMock(return_value=False)
        mock_pdfplumber.open.return_value = mock_pdf

        with patch("pathlib.Path.exists", return_value=True):
            extractor = PDFExtractor(ocr_enabled=False)
            pages = extractor.extract("/fake/doc.pdf")

        assert len(pages[0].tables) == 1
        assert pages[0].tables[0][0][0] == "Analysis"


# ── DOCXExtractor ─────────────────────────────────────────────────────────────

class TestDOCXExtractor:
    def test_raises_on_missing_file(self):
        from ingestion.docx_extractor import DOCXExtractor
        extractor = DOCXExtractor()
        with pytest.raises(FileNotFoundError):
            extractor.extract("/nonexistent/path.docx")

    def test_builds_full_text(self):
        from ingestion.docx_extractor import DOCXExtractor

        # Create a minimal in-memory DOCX
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading("Product Testing Procedure", level=1)
        doc.add_paragraph("This procedure describes the testing of Product X.")

        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            extractor = DOCXExtractor()
            content = extractor.extract(tmp_path)
            assert "Product Testing Procedure" in content.full_text
            assert "Product X" in content.full_text
        finally:
            os.unlink(tmp_path)


# ── Dispatcher ────────────────────────────────────────────────────────────────

class TestDispatcher:
    def test_raises_on_unsupported_type(self, tmp_path):
        (tmp_path / "test.xlsx").touch()
        from ingestion.dispatcher import ingest_document
        with pytest.raises(ValueError, match="Unsupported file type"):
            ingest_document(tmp_path / "test.xlsx")
