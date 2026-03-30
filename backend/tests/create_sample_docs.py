"""
create_sample_docs.py - Generate sample STP and PTP Word documents for testing.

Run with:
    python tests/create_sample_docs.py

Generates:
    tests/sample_stp.docx  - Standard Testing Procedure sample
    tests/sample_ptp.docx  - Product Testing Procedure sample
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def make_stp(output_path: Path) -> None:
    """Generate a realistic-looking STP document."""
    doc = Document()

    # Title
    title = doc.add_heading("Standard Testing Procedure", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("STP-001 | Revision: 3 | Effective Date: 2024-01-15")
    doc.add_paragraph("Prepared by: QC Laboratory | Approved by: QA Manager")
    doc.add_paragraph("")

    # 1. Scope
    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph(
        "This Standard Testing Procedure defines the analytical methods for testing "
        "raw materials and finished products including wheat flour, corn starch, and "
        "related food ingredients."
    )

    # 2. Analyses
    doc.add_heading("2. Analytical Methods", level=1)

    doc.add_heading("2.1 Moisture Content", level=2)
    doc.add_paragraph(
        "Analysis Code: MOISTURE\n"
        "Analysis Type: Physical\n"
        "Common Name: Loss on Drying (LOD)\n"
        "Description: Determination of moisture content by gravimetric method (105°C, 2h)\n"
        "Group: Physical Properties"
    )
    doc.add_paragraph("Components:")
    t1 = doc.add_table(rows=3, cols=5)
    t1.style = "Table Grid"
    headers = ["Parameter", "Num Replicates", "Result Type", "Units", "Typical Range"]
    data = [
        ["Moisture %", "2", "Numeric", "%", "5 - 15"],
    ]
    for i, h in enumerate(headers):
        t1.rows[0].cells[i].text = h
    for row_idx, row in enumerate(data, start=1):
        for col_idx, val in enumerate(row):
            t1.rows[row_idx].cells[col_idx].text = val

    doc.add_heading("2.2 pH Determination", level=2)
    doc.add_paragraph(
        "Analysis Code: PH_VALUE\n"
        "Analysis Type: Chemical\n"
        "Common Name: pH\n"
        "Description: Measurement of acidity/alkalinity using calibrated pH meter\n"
        "Group: Chemical Properties"
    )
    doc.add_paragraph("Components:")
    t2 = doc.add_table(rows=2, cols=5)
    t2.style = "Table Grid"
    for i, h in enumerate(headers):
        t2.rows[0].cells[i].text = h
    ph_data = ["pH Value", "1", "Numeric", "pH units", "4.0 - 7.0"]
    for i, v in enumerate(ph_data):
        t2.rows[1].cells[i].text = v

    doc.add_heading("2.3 Ash Content", level=2)
    doc.add_paragraph(
        "Analysis Code: ASH_CONTENT\n"
        "Analysis Type: Chemical\n"
        "Description: Determination of total ash content by incineration at 550°C\n"
        "Units: % (w/w)\n"
        "Specification: NMT 1.5%"
    )

    doc.add_heading("2.4 Total Plate Count", level=2)
    doc.add_paragraph(
        "Analysis Code: TPC\n"
        "Analysis Type: Microbiological\n"
        "Common Name: Aerobic Plate Count\n"
        "Description: Enumeration of total aerobic mesophilic bacteria\n"
        "Units: CFU/g\n"
        "Specification: NMT 100,000 CFU/g"
    )

    doc.add_heading("2.5 Salmonella", level=2)
    doc.add_paragraph(
        "Analysis Code: SALMONELLA\n"
        "Analysis Type: Microbiological\n"
        "Description: Detection of Salmonella spp. by enrichment and selective plating\n"
        "Result Type: Pass/Fail\n"
        "Specification: Absent in 25g"
    )

    # 3. Specifications Table
    doc.add_heading("3. Specification Limits", level=1)
    spec_table = doc.add_table(rows=7, cols=6)
    spec_table.style = "Table Grid"
    spec_headers = ["Analysis", "Component", "Units", "Spec Rule", "Min", "Max"]
    spec_data = [
        ["MOISTURE",    "Moisture %",    "%",     "Between",    "5.0",   "15.0"],
        ["PH_VALUE",    "pH Value",      "pH",    "Between",    "4.0",   "7.0"],
        ["ASH_CONTENT", "Ash %",         "%",     "LessThan",   "",      "1.5"],
        ["TPC",         "CFU/g",         "CFU/g", "LessThan",   "",      "100000"],
        ["SALMONELLA",  "Result",        "",      "Text",       "",      "Absent"],
    ]
    for i, h in enumerate(spec_headers):
        spec_table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(spec_data, start=1):
        for c_idx, val in enumerate(row):
            spec_table.rows[r_idx].cells[c_idx].text = val

    # 4. Sample Plan
    doc.add_heading("4. Sampling Requirements", level=1)
    doc.add_paragraph(
        "Sample Plan: INCOMING_RM\n"
        "Description: Incoming raw material sampling plan\n"
        "Stage: Receiving\n"
        "Num Samples: 3\n"
        "Sample Size: 200g per sample\n"
        "Test Location: QC Lab"
    )

    doc.save(str(output_path))
    print(f"Created STP sample: {output_path}")


def make_ptp(output_path: Path) -> None:
    """Generate a realistic-looking PTP document."""
    doc = Document()

    title = doc.add_heading("Product Testing Procedure", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("PTP-2024-WF-001 | Wheat Flour Products")
    doc.add_paragraph("Product Group: Baking Ingredients")
    doc.add_paragraph("")

    doc.add_heading("1. Product Information", level=1)
    doc.add_paragraph("Product Name: WHEAT_FLOUR_PLAIN\nDescription: Plain wheat flour for general baking use\nGroup: Flour Products")
    doc.add_paragraph("Product Name: WHEAT_FLOUR_STRONG\nDescription: Strong bread flour, high protein\nGroup: Flour Products")

    doc.add_heading("2. Item Codes", level=1)
    it_table = doc.add_table(rows=3, cols=4)
    it_table.style = "Table Grid"
    for i, h in enumerate(["Item Code", "Description", "Sample Plan", "Site"]):
        it_table.rows[0].cells[i].text = h
    rows = [
        ["WF-001-P", "Plain Wheat Flour 1kg", "INCOMING_RM", "Factory 1"],
        ["WF-002-S", "Strong Bread Flour 1kg", "INCOMING_RM", "Factory 1"],
    ]
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, v in enumerate(row):
            it_table.rows[r_idx].cells[c_idx].text = v

    doc.add_heading("3. Specifications", level=1)
    doc.add_paragraph("Grade: Premium | Stage: Finished Product | Spec Type: Release")

    spec_t = doc.add_table(rows=6, cols=7)
    spec_t.style = "Table Grid"
    for i, h in enumerate(["Product", "Analysis", "Component", "Units", "Min", "Max", "Show on Cert"]):
        spec_t.rows[0].cells[i].text = h
    specs = [
        ["WHEAT_FLOUR_PLAIN", "MOISTURE",    "Moisture %", "%",     "10.0", "14.5", "Y"],
        ["WHEAT_FLOUR_PLAIN", "PH_VALUE",    "pH",         "pH",    "5.5",  "7.0",  "Y"],
        ["WHEAT_FLOUR_PLAIN", "ASH_CONTENT", "Ash %",      "%",     "0.4",  "0.8",  "Y"],
        ["WHEAT_FLOUR_PLAIN", "TPC",         "CFU/g",      "CFU/g", "",     "50000","N"],
        ["WHEAT_FLOUR_PLAIN", "SALMONELLA",  "Result",     "",      "",     "",     "Y"],
    ]
    for r_idx, row in enumerate(specs, start=1):
        for c_idx, v in enumerate(row):
            spec_t.rows[r_idx].cells[c_idx].text = v

    doc.add_heading("4. Sample Plans", level=1)
    doc.add_paragraph(
        "Sample Plan: FINISHED_PRODUCT\n"
        "Description: Finished product release testing\n"
        "Stage: Finished Product\n"
        "Num Samples: 2\n"
        "Quantity: 500g\n"
        "Initial Status: Quarantine"
    )

    doc.save(str(output_path))
    print(f"Created PTP sample: {output_path}")


if __name__ == "__main__":
    out = Path(__file__).parent
    make_stp(out / "sample_stp.docx")
    make_ptp(out / "sample_ptp.docx")
    print("Sample documents created successfully.")
