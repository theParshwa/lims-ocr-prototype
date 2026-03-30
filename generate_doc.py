from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
styles = doc.styles

def style_heading(paragraph, size, color_hex, bold=True, space_before=12, space_after=4):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after  = Pt(space_after)

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_heading(p, 22, '0C1428', space_before=18, space_after=6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A3C6B')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_heading(p, 14, '1A3C6B', space_before=14, space_after=4)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_heading(p, 12, '2563EB', space_before=10, space_after=3)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    p.paragraph_format.left_indent  = Inches(0.3 + 0.2 * level)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_kv(doc, key, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{key}: ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    p.paragraph_format.space_after = Pt(3)
    return p

def shade_row(row, hex_color):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    shade_row(hrow, '0C1428')
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
        shade_row(row, fill)
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.paragraph_format.space_before = Pt(60)
cover.paragraph_format.space_after  = Pt(4)
cover.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('LIMS OCR Document Processor')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x0C, 0x14, 0x28)

sub = doc.add_paragraph()
sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
rs = sub.add_run('AI-Powered Laboratory Document Extraction & Load Sheet Generator')
rs.font.size = Pt(13)
rs.italic = True
rs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)

date_p = doc.add_paragraph()
date_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_before = Pt(16)
dr = date_p.add_run(f'Version 1.0  ·  {datetime.date.today().strftime("%B %Y")}')
dr.font.size = Pt(10)
dr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'Table of Contents')
toc_items = [
    ('1.', 'Executive Summary',                       '3'),
    ('2.', 'Application Overview',                    '3'),
    ('3.', 'Key Features',                            '4'),
    ('4.', 'System Architecture',                     '5'),
    ('5.', 'Technology Stack',                        '6'),
    ('6.', 'AI Agent Pipeline',                       '7'),
    ('7.', 'Correction Capture & RAG Retrieval',      '8'),
    ('8.', 'Supported LIMS Systems',                  '9'),
    ('9.', 'Supported Document Types',                '9'),
    ('10.', 'Data Extraction Sheets',                 '10'),
    ('11.', 'User Interface Walkthrough',             '11'),
    ('12.', 'Deployment',                             '13'),
    ('13.', 'Configuration Reference',                '14'),
    ('14.', 'Security & Data Handling',               '15'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'{num}  {title}')
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6B)
    r2 = p.add_run(f'  {"." * (55 - len(num) - len(title))}  {page}')
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '1. Executive Summary')
add_body(doc,
    'The LIMS OCR Document Processor is an enterprise-grade, AI-powered web application '
    'designed to eliminate the manual effort of extracting data from laboratory documents and '
    'populating LIMS load sheets. It combines state-of-the-art optical character recognition '
    '(OCR), large language model (LLM) extraction, and a retrieval-augmented generation (RAG) '
    'feedback loop to deliver accurate, structured data ready for import into LabWare, '
    'LabVantage, or Veeva Vault LIMS platforms.'
)
add_body(doc,
    'Laboratory teams typically spend hours manually transcribing data from Standard Test '
    'Procedures (STPs), Product Test Plans (PTPs), Specifications, and Methods into spreadsheet '
    'load sheets. This application automates that process end-to-end — upload a PDF, review '
    'the AI-extracted data in an editable grid, correct any errors, and export a fully '
    'formatted Excel load sheet in seconds.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. APPLICATION OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '2. Application Overview')
add_body(doc,
    'The application is structured as a modern full-stack web application with a Python '
    'FastAPI backend and a React TypeScript frontend. Users interact with the app entirely '
    'through a browser — no desktop software installation is required.'
)

add_h2(doc, 'Problem Statement')
add_body(doc,
    'Laboratory information management systems require large volumes of structured data to be '
    'loaded from paper or PDF-based source documents. This data entry is:')
add_bullet(doc, 'Time-consuming — analysts spend hours copying tables and fields manually')
add_bullet(doc, 'Error-prone — transcription mistakes cause failed LIMS imports or incorrect test setups')
add_bullet(doc, 'Not scalable — load sheet creation becomes a bottleneck during product launches or audits')

add_h2(doc, 'Solution')
add_body(doc,
    'The LIMS OCR Document Processor addresses these pain points by:')
add_bullet(doc, 'Automatically extracting text from PDFs using pdfplumber (digital) or Tesseract OCR (scanned)')
add_bullet(doc, 'Classifying the document type (STP, PTP, SPEC, METHOD, SOP) using an AI agent')
add_bullet(doc, 'Extracting 30 categories of LIMS-relevant data using GPT-4o with LIMS-specific prompts')
add_bullet(doc, 'Presenting extracted data in editable grids for human review and correction')
add_bullet(doc, 'Learning from corrections via RAG to improve future extractions automatically')
add_bullet(doc, 'Exporting structured Excel load sheets formatted for the selected LIMS system')

# ══════════════════════════════════════════════════════════════════════════════
#  3. KEY FEATURES
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '3. Key Features')

features = [
    ('Multi-LIMS Support',
     'Pre-configured extraction and export templates for LabWare LIMS, LabVantage, and Veeva Vault. '
     'Selecting a LIMS system at startup tailors all prompts, field names, and export formats.'),
    ('Document Type Awareness',
     'Users select the document type at upload time (STP, PTP, SPEC, METHOD, SOP, OTHER). '
     'The AI agent adjusts its extraction focus accordingly — e.g., STPs prioritise instrument '
     'and method parameters, while SPECs focus on acceptance limits and reference ranges.'),
    ('30-Sheet Extraction',
     'The system extracts data into 30 structured categories matching a full LIMS load sheet: '
     'methods, instruments, tests, specifications, stability, reagents, audit fields, and more.'),
    ('Confidence Scoring',
     'Every extracted field is given a confidence score (0–1). Fields below 0.6 are highlighted '
     'amber for review; validation errors are highlighted red. A colour-coded confidence bar '
     'is shown for each sheet.'),
    ('Editable Data Grids',
     'Extracted data is presented in AG Grid tables. Users can edit any cell inline before '
     'exporting, ensuring full human oversight.'),
    ('Correction Capture + RAG',
     'When a user saves corrections, the system automatically diffs the original vs corrected '
     'values, stores each correction as a training example, and generates vector embeddings. '
     'Future extractions retrieve the most relevant past corrections via cosine similarity.'),
    ('Configurable Agent Prompts',
     'All four AI agent prompts (Classifier, Extractor, Mapper, Validator) are fully editable '
     'from the Agent tab. Users can view rendered previews with placeholders filled.'),
    ('Custom Load Sheet Templates',
     'In the Configure tab, users can paste their organisation\'s actual LIMS column headers. '
     'These are injected directly into agent prompts to ensure extracted field names match '
     'the target system exactly.'),
    ('Excel Export',
     'One-click export generates a multi-tab Excel workbook (.xlsx) with each extraction '
     'sheet on its own tab, ready to import into the LIMS.'),
    ('Job History & Dashboard',
     'Every uploaded document creates a job record. The dashboard shows status (processing / '
     'complete / failed), document type, confidence, and provides access to re-processing.'),
]

for title, desc in features:
    add_h3(doc, f'✦  {title}')
    add_body(doc, desc)

# ══════════════════════════════════════════════════════════════════════════════
#  4. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '4. System Architecture')
add_body(doc,
    'The application follows a clean client-server architecture with a clear separation of '
    'concerns between the React frontend, FastAPI backend, and external AI/storage services.'
)

add_h2(doc, 'High-Level Architecture')
arch_rows = [
    ('Browser (React SPA)', 'User interface — file upload, data review, configuration, export'),
    ('FastAPI Backend',      'REST API, background processing, business logic, job management'),
    ('Extraction Pipeline',  '4-pass LLM extraction orchestrator with chunking and confidence scoring'),
    ('RAG Store (SQLite)',   'Vector embeddings + correction history for retrieval-augmented generation'),
    ('OpenAI / Claude',      'GPT-4o (default) or Claude for document classification and extraction'),
    ('Tesseract OCR',        'Scanned PDF text extraction via pdf2image + pytesseract'),
    ('Excel Export',         'openpyxl-based multi-sheet workbook generation'),
]
add_table(doc,
    ['Component', 'Responsibility'],
    arch_rows,
    col_widths=[2.2, 4.3]
)

add_h2(doc, 'Request Flow')
steps = [
    'User selects LIMS system on first launch → stored in localStorage and backend config',
    'User uploads PDF with document type hint via drag-and-drop zone',
    'Backend receives file → creates job record → starts background processing task',
    'Pipeline extracts text (pdfplumber or Tesseract), chunks into 4000-char segments',
    'RAG retriever queries embedding store for top-5 similar past corrections/examples',
    'Classifier agent identifies document type (overridden by user hint if provided)',
    '4-pass LLM extraction: foundational data → components → specifications → plans/codes',
    'Results stored in SQLite with confidence scores → job status set to complete',
    'Frontend polls job status → renders 30-tab AG Grid with colour-coded confidence',
    'User reviews and edits → saves → corrections captured, embedded, stored in RAG',
    'User exports → Excel workbook generated and downloaded',
]
for i, step in enumerate(steps, 1):
    add_bullet(doc, f'Step {i}: {step}')

# ══════════════════════════════════════════════════════════════════════════════
#  5. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '5. Technology Stack')

add_h2(doc, 'Backend')
be_rows = [
    ('Python 3.11',           'Core language'),
    ('FastAPI',               'Async REST API framework with automatic OpenAPI docs'),
    ('LangChain',             'LLM orchestration, prompt management, chain composition'),
    ('OpenAI SDK',            'GPT-4o for classification, extraction, mapping, validation'),
    ('pdfplumber',            'High-fidelity text and table extraction from digital PDFs'),
    ('pytesseract + pdf2image','OCR pipeline for scanned/image-based PDFs'),
    ('openpyxl',              'Excel (.xlsx) generation for load sheet export'),
    ('SQLite + aiosqlite',    'Async embedded database for job history and RAG store'),
    ('pydantic-settings',     'Type-safe configuration management via .env files'),
    ('NumPy',                 'Vector operations for cosine similarity in RAG retrieval'),
]
add_table(doc, ['Library / Framework', 'Purpose'], be_rows, col_widths=[2.2, 4.3])

add_h2(doc, 'Frontend')
fe_rows = [
    ('React 18 + TypeScript', 'Component-based SPA with full type safety'),
    ('Vite',                  'Lightning-fast build tool and dev server'),
    ('Tailwind CSS',          'Utility-first CSS with custom enterprise dark palette'),
    ('AG Grid Community',     'High-performance editable data grid for extraction results'),
    ('Axios',                 'HTTP client for REST API communication'),
    ('Lucide React',          'Consistent icon library'),
    ('Heroicons',             'Additional UI icons'),
]
add_table(doc, ['Library / Framework', 'Purpose'], fe_rows, col_widths=[2.2, 4.3])

add_h2(doc, 'AI & ML')
ai_rows = [
    ('OpenAI GPT-4o',               'LLM for all extraction tasks'),
    ('text-embedding-3-small',      'OpenAI embedding model for RAG vector store'),
    ('TF-IDF (fallback)',           'Local bag-of-words embedding when OpenAI unavailable'),
    ('Cosine Similarity',           'RAG retrieval scoring for finding relevant past examples'),
]
add_table(doc, ['Technology', 'Role'], ai_rows, col_widths=[2.5, 4.0])

add_h2(doc, 'Infrastructure')
infra_rows = [
    ('Vercel',   'Frontend hosting (free tier, global CDN, auto-deploy from GitHub)'),
    ('Render',   'Backend hosting (free tier, Docker-based deployment)'),
    ('Supabase', 'Optional managed PostgreSQL (free tier, replaces SQLite in production)'),
    ('GitHub',   'Source control and CI/CD trigger for auto-deployments'),
    ('Docker',   'Backend containerisation for consistent deployments'),
]
add_table(doc, ['Service', 'Role'], infra_rows, col_widths=[1.5, 5.0])

# ══════════════════════════════════════════════════════════════════════════════
#  6. AI AGENT PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '6. AI Agent Pipeline')
add_body(doc,
    'The application uses a multi-agent pipeline with four specialised AI agents, each with '
    'a dedicated, editable system prompt. All prompts are LIMS-aware and document-type-aware.'
)

agents = [
    ('Classifier Agent',
     'Analyses the raw document text and determines the document type (STP, PTP, SPEC, '
     'METHOD, SOP, OTHER). The result is used to configure all downstream agents. If the '
     'user provided a document type hint at upload time, this overrides the classifier output.'),
    ('Extractor Agent',
     'The primary workhorse. Processes the document in 4000-character chunks with 400-character '
     'overlap to handle large documents. Uses a 4-pass extraction strategy:\n'
     '  Pass 1 (Foundational): Document ID, title, version, effective date, product details\n'
     '  Pass 2 (Components): Methods, instruments, reagents, equipment, personnel\n'
     '  Pass 3 (Specifications): Acceptance limits, reference ranges, test parameters\n'
     '  Pass 4 (Plans & Codes): Test plans, audit fields, change control references\n'
     'Each field receives a confidence score (0.0–1.0) based on extraction certainty.'),
    ('Mapper Agent',
     'Maps extracted field names and values to the target LIMS system\'s exact field names '
     'and controlled vocabulary. Uses the user-supplied load sheet template (column headers) '
     'to ensure perfect alignment with the organisation\'s LIMS configuration.'),
    ('Validator Agent',
     'Validates the mapped data against LIMS business rules: required field checks, data '
     'type validation, controlled vocabulary lookups, and cross-field consistency checks. '
     'Flags validation errors for human review.'),
]
for name, desc in agents:
    add_h3(doc, name)
    add_body(doc, desc)

add_h2(doc, 'Prompt Customisation')
add_body(doc,
    'All four agent prompts are fully editable from the Agent tab in the UI. Each prompt '
    'supports the following dynamic placeholders:'
)
add_kv(doc, '{lims_system}',          'Replaced with "LabWare", "LabVantage", or "Veeva Vault"')
add_kv(doc, '{document_type}',        'Replaced with "STP", "PTP", "SPEC", "METHOD", "SOP", or "OTHER"')
add_kv(doc, '{load_sheet_template}',  'Replaced with the column headers pasted in Configure → Load Sheet')
add_kv(doc, '{training_context}',     'Replaced with relevant RAG-retrieved correction examples')

# ══════════════════════════════════════════════════════════════════════════════
#  7. CORRECTION CAPTURE & RAG
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '7. Correction Capture & RAG Retrieval')
add_body(doc,
    'The RAG (Retrieval-Augmented Generation) system allows the application to learn from '
    'user corrections over time, continuously improving extraction accuracy without requiring '
    'model retraining.'
)

add_h2(doc, 'How It Works')
rag_steps = [
    ('Correction Detection',
     'When a user edits extracted data and clicks Save, the system diffs the original AI '
     'output against the user\'s corrections, identifying every changed field.'),
    ('Correction Storage',
     'Each changed field is stored as a CorrectionExample record in SQLite, capturing: '
     'the sheet name, field name, original AI value, corrected value, surrounding context '
     'text, and document type.'),
    ('Embedding Generation',
     'An embedding vector is generated for each correction using OpenAI text-embedding-3-small '
     '(or a TF-IDF fallback). The embed text includes sheet, field, context, and document type '
     'to enable semantic matching.'),
    ('RAG Retrieval at Extraction Time',
     'When a new document is uploaded, the system reads the first 8KB of text, generates '
     'an embedding, and retrieves the top-5 most similar past corrections from the vector '
     'store using cosine similarity.'),
    ('Few-Shot Injection',
     'Retrieved corrections are formatted as few-shot examples and injected into the '
     'Extractor agent\'s prompt as {training_context}. This guides the LLM to make the '
     'same corrections for similar documents in future.'),
]
for i, (title, desc) in enumerate(rag_steps, 1):
    add_h3(doc, f'{i}. {title}')
    add_body(doc, desc)

add_h2(doc, 'Technical Implementation')
rag_tech = [
    ('Embedding Model',  'OpenAI text-embedding-3-small (1536-dim) with disk cache (MD5-keyed JSON)'),
    ('Fallback Embedder','TF-IDF bag-of-words (512-term vocabulary) — no API key required'),
    ('Similarity Metric','Cosine similarity with vector length normalisation'),
    ('Storage',          'RAGEmbedding table in SQLite — stores source type, source ID, text, JSON embedding'),
    ('Retrieval',        'Loads all embeddings, scores in memory, returns top-k (default k=5)'),
    ('Sources',          '"training" (manual examples) and "correction" (auto-captured) — corrections ranked first'),
]
add_table(doc, ['Aspect', 'Detail'], rag_tech, col_widths=[1.8, 4.7])

# ══════════════════════════════════════════════════════════════════════════════
#  8. SUPPORTED LIMS SYSTEMS
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '8. Supported LIMS Systems')
lims_rows = [
    ('LabWare LIMS',
     'Industry-standard LIMS used in pharmaceutical, food, and environmental labs. '
     'Extraction targets LabWare\'s entity model: Methods, Instruments, Tests, Specifications.'),
    ('LabVantage',
     'Enterprise LIMS with a configurable data model. Extraction aligns to LabVantage '
     'SampleLogin, WorkSheets, and Results entities.'),
    ('Veeva Vault',
     'Cloud-based LIMS used primarily in regulated pharmaceutical and biotech environments. '
     'Extraction targets Veeva\'s document-centric model with audit trail fields.'),
]
add_table(doc, ['LIMS System', 'Description'], lims_rows, col_widths=[1.8, 4.7])

add_body(doc,
    'Selecting a LIMS system at startup configures: agent prompt phrasing, field name '
    'vocabulary, controlled value lists, and the Excel export column layout. The selection '
    'is stored persistently in both the browser (localStorage) and the backend configuration.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  9. SUPPORTED DOCUMENT TYPES
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '9. Supported Document Types')
doc_rows = [
    ('STP', 'Standard Test Procedure',
     'Prioritises instrument parameters, method steps, reagent preparation, and equipment calibration'),
    ('PTP', 'Product Test Plan',
     'Focuses on test sequences, acceptance criteria, sampling plans, and product specifications'),
    ('SPEC', 'Specification',
     'Targets acceptance limits, reference ranges, units of measure, and out-of-specification handling'),
    ('METHOD', 'Analytical Method',
     'Extracts method parameters, system suitability, column conditions, and reagent details'),
    ('SOP', 'Standard Operating Procedure',
     'Captures procedural steps, roles/responsibilities, equipment, and quality review fields'),
    ('OTHER', 'Other / Unknown',
     'General extraction without document-type-specific focus; all fields attempted'),
]
add_table(doc,
    ['Code', 'Full Name', 'Extraction Focus'],
    doc_rows,
    col_widths=[0.7, 1.8, 4.0]
)

# ══════════════════════════════════════════════════════════════════════════════
#  10. DATA EXTRACTION SHEETS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '10. Data Extraction Sheets')
add_body(doc,
    'Extracted data is organised into 30 sheets, each corresponding to a tab in the '
    'data preview grid and a sheet in the exported Excel workbook.'
)

sheets = [
    ('01', 'Cover',                'Document metadata, title, version, effective date, approvals'),
    ('02', 'Document Control',     'Change history, revision log, document control numbers'),
    ('03', 'Product Info',         'Product name, code, formulation, regulatory details'),
    ('04', 'Test Overview',        'Test summary, scope, purpose, applicable standards'),
    ('05', 'Instruments',          'Equipment IDs, model numbers, calibration requirements'),
    ('06', 'Reagents',             'Reagent names, grades, supplier, preparation instructions'),
    ('07', 'Standards',            'Reference standards, CAS numbers, purity, storage'),
    ('08', 'Sample Prep',          'Sample preparation steps, weights, dilution factors'),
    ('09', 'System Suitability',   'SST parameters, limits, frequency, acceptance criteria'),
    ('10', 'Method Parameters',    'Column conditions, mobile phase, flow rate, detection wavelength'),
    ('11', 'Calculations',         'Formulae, units, rounding rules, significant figures'),
    ('12', 'Specifications',       'Acceptance limits, reference ranges, out-of-spec handling'),
    ('13', 'Test Methods',         'Compendial method references, in-house method IDs'),
    ('14', 'Sampling Plan',        'Sampling frequency, sample size, container selection'),
    ('15', 'Stability',            'Storage conditions, time points, degradation pathways'),
    ('16', 'Packaging',            'Container closure, labelling requirements, storage'),
    ('17', 'Microbiology',         'Microbial limits, sterility tests, environmental monitoring'),
    ('18', 'Impurities',           'Impurity thresholds, identification/qualification limits'),
    ('19', 'Related Substances',   'Known related substances, ICH classification, limits'),
    ('20', 'Dissolution',          'Dissolution media, apparatus, speed, time points, limits'),
    ('21', 'Content Uniformity',   'CU method, acceptance value, individual limits'),
    ('22', 'Assay',                'Assay method, range, acceptance criteria, reference standard'),
    ('23', 'Water Content',        'KF method, LOD conditions, limits'),
    ('24', 'Physical Tests',       'Appearance, description, colour, pH, viscosity'),
    ('25', 'Audit Fields',         'Reviewer, approver, effective date, supersedes version'),
    ('26', 'Change Control',       'Change control reference, justification, impact assessment'),
    ('27', 'Training Records',     'Training requirements, personnel qualifications'),
    ('28', 'Deviations',           'Deviation reference fields, impact, CAPA linkage'),
    ('29', 'References',           'Cited documents, pharmacopoeia references, ICH guidelines'),
    ('30', 'Attachments',          'Supplementary tables, chromatograms, figures referenced'),
]
add_table(doc,
    ['#', 'Sheet Name', 'Key Data Captured'],
    sheets,
    col_widths=[0.4, 1.7, 4.4]
)

# ══════════════════════════════════════════════════════════════════════════════
#  11. UI WALKTHROUGH
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '11. User Interface Walkthrough')

screens = [
    ('LIMS Selector (Onboarding)',
     'On first launch, users see three cards for LabWare, LabVantage, and Veeva Vault. '
     'Selecting a LIMS configures the entire application. The selection is persisted and '
     'can be changed anytime from Configure → LIMS System.'),
    ('Upload Tab',
     'Drag-and-drop zone supporting PDF files. Before dropping, users select the document '
     'type from a dropdown (STP, PTP, SPEC, METHOD, SOP, OTHER). A progress stepper shows '
     'real-time extraction stages: Upload → OCR → Classify → Extract → Map → Validate. '
     'Multiple files can be queued.'),
    ('Dashboard Tab',
     'A tabular view of all processed documents showing: status indicator (green/amber/red '
     'dot), filename, document type chip, processing status badge, date/time, job ID, and '
     'action buttons. Clicking a completed row opens the data preview panel.'),
    ('Data Preview (Sheet Grids)',
     '30 tabbed AG Grid tables with inline editing. Each tab shows a confidence bar '
     '(green ≥0.8, amber ≥0.6, red <0.6). Individual cells are colour-coded: yellow for '
     'low-confidence fields, red for validation errors. Export button generates the Excel '
     'workbook. Save button captures corrections into the RAG store.'),
    ('Agent Tab',
     'Two sub-tabs: Agent Prompts and Training Examples. '
     'Agent Prompts shows all four agent prompts (Classifier, Extractor, Mapper, Validator) '
     'with expand/collapse, Edit Template / Preview (rendered) toggle, and Save/Reset per agent. '
     'Training Examples allows manual addition of field-level training examples.'),
    ('Configure Tab',
     'Four configuration sections:\n'
     '  LIMS System: Change the active LIMS (LabWare / LabVantage / Veeva)\n'
     '  Load Sheet: Paste column headers per LIMS to align extraction field names\n'
     '  Output Sheets: Checkbox grid to enable/disable which of the 30 sheets are extracted\n'
     '  AI Settings: Temperature, max tokens, chunk size, chunk overlap, confidence threshold'),
]
for title, desc in screens:
    add_h2(doc, title)
    add_body(doc, desc)

# ══════════════════════════════════════════════════════════════════════════════
#  12. DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_h1(doc, '12. Deployment')

add_h2(doc, 'Architecture')
add_body(doc,
    'The application is deployed as two separate services connected via environment variables:'
)
add_bullet(doc, 'Frontend → Vercel (global CDN, auto-deploy from GitHub, free tier)')
add_bullet(doc, 'Backend → Render (Docker-based, free tier, spins up on request)')
add_bullet(doc, 'Database → SQLite (embedded, file-based) or Supabase PostgreSQL (managed, free tier)')

add_h2(doc, 'Environment Variables')
env_rows = [
    ('OPENAI_API_KEY',    'Backend', 'Required. OpenAI API key for LLM extraction and embeddings'),
    ('DATABASE_URL',      'Backend', 'Optional. PostgreSQL URL for Supabase. Defaults to SQLite'),
    ('VITE_API_URL',      'Frontend', 'Required. Full URL of the deployed backend API'),
    ('CORS_ORIGINS',      'Backend', 'Required. Comma-separated list of allowed frontend origins'),
]
add_table(doc,
    ['Variable', 'Service', 'Description'],
    env_rows,
    col_widths=[2.0, 1.0, 3.5]
)

add_h2(doc, 'Local Development')
add_body(doc, 'Prerequisites: Python 3.11+, Node.js 18+, Tesseract OCR (system package)')
add_kv(doc, 'Start Backend',  'cd backend && py -m uvicorn api.main:app --reload --port 8000')
add_kv(doc, 'Start Frontend', 'cd frontend && npm install && npm run dev')
add_kv(doc, 'Backend URL',    'http://localhost:8000')
add_kv(doc, 'Frontend URL',   'http://localhost:5173')

# ══════════════════════════════════════════════════════════════════════════════
#  13. CONFIGURATION REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '13. Configuration Reference')

add_h2(doc, 'backend/.env')
config_rows = [
    ('OPENAI_API_KEY',      'sk-...',      'OpenAI API key'),
    ('MODEL_NAME',          'gpt-4o',       'Specific model to use'),
    ('CHUNK_SIZE',          '4000',         'Characters per extraction chunk'),
    ('CHUNK_OVERLAP',       '400',          'Overlap between chunks'),
    ('CONFIDENCE_THRESHOLD','0.6',          'Minimum confidence for auto-accept'),
    ('MAX_TOKENS',          '4096',         'Maximum LLM response tokens'),
    ('TEMPERATURE',         '0.1',          'LLM temperature (lower = more deterministic)'),
    ('DATABASE_URL',        '(sqlite)',      'Database connection string'),
    ('CORS_ORIGINS',        'http://localhost:5173', 'Allowed CORS origins'),
]
add_table(doc,
    ['Setting', 'Default', 'Description'],
    config_rows,
    col_widths=[2.0, 1.4, 3.1]
)

add_h2(doc, 'Runtime Configuration (Configure Tab)')
add_body(doc,
    'The following settings are configurable at runtime without restarting the server, '
    'via the Configure tab in the UI or the GET/PUT /api/config REST endpoint:'
)
add_bullet(doc, 'Active LIMS system (LabWare / LabVantage / Veeva Vault)')
add_bullet(doc, 'Enabled output sheets (which of the 30 sheets to extract)')
add_bullet(doc, 'Load sheet templates (per LIMS column header definitions)')
add_bullet(doc, 'AI temperature, max tokens, chunk size, confidence threshold')
add_bullet(doc, 'Document type list and display names')

# ══════════════════════════════════════════════════════════════════════════════
#  14. SECURITY & DATA HANDLING
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, '14. Security & Data Handling')

add_h2(doc, 'Data Privacy')
add_body(doc,
    'Uploaded documents are processed server-side and stored temporarily in the backend\'s '
    'upload directory. Document text is sent to the OpenAI API for extraction. '
    'Organisations handling sensitive documents should review OpenAI\'s data '
    'handling policies and consider using a self-hosted LLM for on-premises deployment.'
)

add_h2(doc, 'API Key Security')
add_bullet(doc, 'API keys are stored in .env files, never committed to source control')
add_bullet(doc, '.gitignore excludes all .env files and the SQLite database file')
add_bullet(doc, 'Render and Vercel environment variables are encrypted at rest')
add_bullet(doc, 'Keys are never exposed to the frontend — all AI calls are server-side')

add_h2(doc, 'Input Validation')
add_bullet(doc, 'File uploads validated for PDF MIME type and size limits')
add_bullet(doc, 'All API inputs validated via Pydantic models with strict typing')
add_bullet(doc, 'CORS configured to allow only the registered frontend origin')
add_bullet(doc, 'SQL injection not applicable — SQLAlchemy ORM with parameterised queries')

add_h2(doc, 'Audit Trail')
add_bullet(doc, 'All extraction jobs stored with timestamp, status, and result hash')
add_bullet(doc, 'Correction history preserved — original AI output is never overwritten')
add_bullet(doc, 'RAG correction store provides complete audit trail of AI improvement over time')

# ══════════════════════════════════════════════════════════════════════════════
#  FOOTER on each page
# ══════════════════════════════════════════════════════════════════════════════
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run('LIMS OCR Document Processor  ·  Confidential  ·  Internal Use Only')
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = r'd:\College\LIMS OCR\LIMS_OCR_Application_Overview.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
